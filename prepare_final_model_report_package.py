from pathlib import Path
import json
import math

import pandas as pd
import matplotlib.pyplot as plt

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError as error:
    raise ImportError("Missing package: python-docx. Install it using: pip install python-docx") from error


OUTPUT_DIR = Path("outputs") / "reports" / "final_model_report_package_strict_detailed"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

FINAL_MODEL_DIR = Path("outputs") / "models" / "final_7class_model"
INTERNAL_NEIGHBOR_DIR = Path("outputs") / "models" / "visium_7class_neighborhood_classification"
STABLE_TOP20_DIR = Path("outputs") / "models" / "visium_stable_gene_classification_top20"
EXTERNAL_DIR = Path("outputs") / "models" / "external_spatialDLPFC_validation"

LABEL_ORDER = [
    "Layer1",
    "Layer2",
    "Layer3",
    "Layer4",
    "Layer5",
    "Layer6",
    "WM",
]

EXPECTED_EXTERNAL_SAMPLES = [
    "Br6522_ant",
    "Br6522_mid",
    "Br8667_post",
]


def require_file(path):
    path = Path(path)

    if not path.exists():
        raise FileNotFoundError(f"Required file not found: {path}")

    return path


def read_csv(path):
    return pd.read_csv(require_file(path))


def read_csv_indexed(path):
    return pd.read_csv(require_file(path), index_col=0)


def read_json(path):
    with open(require_file(path), "r", encoding="utf-8") as file:
        return json.load(file)


def safe_float(value):
    try:
        return float(value)
    except Exception:
        return float("nan")


def format_number(value, digits=4):
    if value is None:
        return ""

    if isinstance(value, float) and math.isnan(value):
        return ""

    try:
        return f"{float(value):.{digits}f}"
    except Exception:
        return str(value)


def format_count(value):
    try:
        if isinstance(value, float) and math.isnan(value):
            return ""
        return f"{int(float(value)):,}"
    except Exception:
        return str(value)


def format_cell(value):
    if isinstance(value, float):
        return format_number(value)

    if isinstance(value, int):
        return format_count(value)

    return str(value)


def difference(new_value, old_value):
    return safe_float(new_value) - safe_float(old_value)


def relative_drop(new_value, old_value):
    old_value = safe_float(old_value)
    new_value = safe_float(new_value)

    if old_value == 0 or math.isnan(old_value) or math.isnan(new_value):
        return float("nan")

    return (new_value - old_value) / old_value


def select_row(df, filters, fallback_sort_column=None):
    selected = df.copy()

    for column, value in filters.items():
        if column not in selected.columns:
            raise ValueError(f"Column '{column}' not found while selecting row with filters {filters}")

        selected = selected[selected[column].astype(str) == str(value)]

    if len(selected) == 0:
        if fallback_sort_column and fallback_sort_column in df.columns:
            return df.sort_values(fallback_sort_column, ascending=False).iloc[0]

        raise ValueError(f"No row matched filters: {filters}")

    return selected.iloc[0]


def normalize_classification_report(df):
    df = df.copy()

    if "label" not in df.columns:
        first_column = df.columns[0]
        df = df.rename(columns={first_column: "label"})

    df["label"] = df["label"].astype(str)
    df = df[df["label"].isin(LABEL_ORDER)].copy()

    if "f1-score" in df.columns:
        df["f1-score"] = df["f1-score"].astype(float)

    if "precision" in df.columns:
        df["precision"] = df["precision"].astype(float)

    if "recall" in df.columns:
        df["recall"] = df["recall"].astype(float)

    if "support" in df.columns:
        df["support"] = df["support"].astype(float)

    df["label"] = pd.Categorical(df["label"], categories=LABEL_ORDER, ordered=True)
    df = df.sort_values("label").reset_index(drop=True)
    df["label"] = df["label"].astype(str)

    return df


def read_classification_report(path):
    return normalize_classification_report(read_csv(path))


def read_confusion_matrix(path):
    matrix = read_csv_indexed(path)

    matrix.index = matrix.index.astype(str)
    matrix.columns = matrix.columns.astype(str)

    matrix = matrix.reindex(index=LABEL_ORDER, columns=LABEL_ORDER, fill_value=0)
    matrix = matrix.fillna(0).astype(int)

    return matrix


def dataframe_to_markdown(df):
    df = df.copy()

    headers = [str(column) for column in df.columns]
    rows = []

    for _, row in df.iterrows():
        rows.append([format_cell(row[column]) for column in df.columns])

    widths = []

    for column_index, header in enumerate(headers):
        max_width = len(header)

        for row in rows:
            max_width = max(max_width, len(str(row[column_index])))

        widths.append(max_width)

    header_line = "| " + " | ".join(headers[index].ljust(widths[index]) for index in range(len(headers))) + " |"
    separator_line = "| " + " | ".join("-" * widths[index] for index in range(len(headers))) + " |"

    row_lines = []

    for row in rows:
        row_lines.append("| " + " | ".join(str(row[index]).ljust(widths[index]) for index in range(len(row))) + " |")

    return "\n".join([header_line, separator_line] + row_lines)


def load_inputs():
    metadata = read_json(FINAL_MODEL_DIR / "final_model_metadata.json")

    selected_genes = read_csv(FINAL_MODEL_DIR / "final_selected_stable_genes.csv")

    final_training_audit = read_csv(FINAL_MODEL_DIR / "final_internal_training_feature_audit.csv")

    final_training_label_counts = read_csv(FINAL_MODEL_DIR / "final_internal_training_label_counts.csv")

    internal_summary = read_csv(
        INTERNAL_NEIGHBOR_DIR / "neighborhood_7class_combined_model_comparison_summary.csv"
    )

    internal_raw = select_row(
        internal_summary,
        {
            "feature_mode": "stable_genes",
            "neighbor_k": 10,
            "model": "linear_svm",
            "prediction_type": "raw",
        },
    )

    internal_smoothed = select_row(
        internal_summary,
        {
            "feature_mode": "stable_genes",
            "neighbor_k": 10,
            "model": "linear_svm",
            "prediction_type": "smoothed",
        },
    )

    stable_top20_summary = read_csv(
        STABLE_TOP20_DIR / "ml_model_comparison_summary.csv"
    )

    stable_top20_baseline = select_row(
        stable_top20_summary,
        {
            "model": "linear_svm",
        },
        fallback_sort_column="weighted_f1",
    )

    external_summary = read_csv(
        EXTERNAL_DIR / "external_spatialDLPFC_validation_summary.csv"
    )

    external_raw = select_row(
        external_summary,
        {
            "sample_id": "ALL_EXTERNAL_LABELLED",
            "prediction_type": "raw",
        },
    )

    external_smoothed = select_row(
        external_summary,
        {
            "sample_id": "ALL_EXTERNAL_LABELLED",
            "prediction_type": "smoothed",
        },
    )

    internal_raw_classwise = read_classification_report(
        INTERNAL_NEIGHBOR_DIR / "neighborhood_classification_report_stable_genes_k10_linear_svm_raw.csv"
    )

    internal_smoothed_classwise = read_classification_report(
        INTERNAL_NEIGHBOR_DIR / "neighborhood_classification_report_stable_genes_k10_linear_svm_smoothed.csv"
    )

    external_classwise_all = read_csv(
        EXTERNAL_DIR / "external_spatialDLPFC_classwise_reports_combined.csv"
    )

    external_classwise_all["sample_id"] = external_classwise_all["sample_id"].astype(str)
    external_classwise_all["prediction_type"] = external_classwise_all["prediction_type"].astype(str)
    external_classwise_all["label"] = external_classwise_all["label"].astype(str)

    external_raw_classwise = normalize_classification_report(
        external_classwise_all[
            (external_classwise_all["sample_id"] == "ALL_EXTERNAL_LABELLED")
            & (external_classwise_all["prediction_type"] == "raw")
            & (external_classwise_all["label"].isin(LABEL_ORDER))
        ].copy()
    )

    external_smoothed_classwise = normalize_classification_report(
        external_classwise_all[
            (external_classwise_all["sample_id"] == "ALL_EXTERNAL_LABELLED")
            & (external_classwise_all["prediction_type"] == "smoothed")
            & (external_classwise_all["label"].isin(LABEL_ORDER))
        ].copy()
    )

    external_sample_summary = external_summary[
        external_summary["sample_id"].astype(str).isin(EXPECTED_EXTERNAL_SAMPLES)
    ].copy()

    external_feature_audit = read_csv(
        EXTERNAL_DIR / "external_feature_availability_audit.csv"
    )

    internal_confusion_raw = read_confusion_matrix(
        INTERNAL_NEIGHBOR_DIR / "confusion_matrix_stable_genes_k10_linear_svm_raw.csv"
    )

    internal_confusion_smoothed = read_confusion_matrix(
        INTERNAL_NEIGHBOR_DIR / "confusion_matrix_stable_genes_k10_linear_svm_smoothed.csv"
    )

    external_confusion_raw = read_confusion_matrix(
        EXTERNAL_DIR / "confusion_matrix_external_ALL_EXTERNAL_LABELLED_raw.csv"
    )

    external_confusion_smoothed = read_confusion_matrix(
        EXTERNAL_DIR / "confusion_matrix_external_ALL_EXTERNAL_LABELLED_smoothed.csv"
    )

    return {
        "metadata": metadata,
        "selected_genes": selected_genes,
        "final_training_audit": final_training_audit,
        "final_training_label_counts": final_training_label_counts,
        "stable_top20_baseline": stable_top20_baseline,
        "internal_raw": internal_raw,
        "internal_smoothed": internal_smoothed,
        "external_raw": external_raw,
        "external_smoothed": external_smoothed,
        "internal_raw_classwise": internal_raw_classwise,
        "internal_smoothed_classwise": internal_smoothed_classwise,
        "external_raw_classwise": external_raw_classwise,
        "external_smoothed_classwise": external_smoothed_classwise,
        "external_classwise_all": external_classwise_all,
        "external_sample_summary": external_sample_summary,
        "external_feature_audit": external_feature_audit,
        "internal_confusion_raw": internal_confusion_raw,
        "internal_confusion_smoothed": internal_confusion_smoothed,
        "external_confusion_raw": external_confusion_raw,
        "external_confusion_smoothed": external_confusion_smoothed,
    }


def get_major_confusions(confusion_matrix, top_n=2):
    rows = []

    for true_label in LABEL_ORDER:
        true_total = int(confusion_matrix.loc[true_label].sum())
        correct = int(confusion_matrix.loc[true_label, true_label])

        off_diagonal = confusion_matrix.loc[true_label].copy()
        off_diagonal[true_label] = 0
        off_diagonal = off_diagonal.sort_values(ascending=False)

        confusions = []

        for predicted_label, count in off_diagonal.head(top_n).items():
            count = int(count)

            if count <= 0:
                continue

            rate = count / true_total if true_total else 0
            confusions.append(f"{predicted_label}: {count} ({rate:.1%})")

        rows.append(
            {
                "True layer": true_label,
                "True support": true_total,
                "Correct predictions": correct,
                "Correct rate": correct / true_total if true_total else float("nan"),
                "Main wrong predictions": "; ".join(confusions) if confusions else "None",
            }
        )

    return pd.DataFrame(rows)


def build_external_label_support_table(external_classwise_all):
    smoothed = external_classwise_all[
        (external_classwise_all["prediction_type"] == "smoothed")
        & (external_classwise_all["sample_id"].isin(EXPECTED_EXTERNAL_SAMPLES))
        & (external_classwise_all["label"].isin(LABEL_ORDER))
    ].copy()

    table = smoothed.pivot_table(
        index="sample_id",
        columns="label",
        values="support",
        aggfunc="first",
        fill_value=0,
    )

    table = table.reindex(index=EXPECTED_EXTERNAL_SAMPLES, columns=LABEL_ORDER, fill_value=0)
    table = table.reset_index()
    table.columns.name = None

    return table


def build_summary_tables(data):
    metadata = data["metadata"]
    baseline = data["stable_top20_baseline"]
    internal_raw = data["internal_raw"]
    internal_smoothed = data["internal_smoothed"]
    external_raw = data["external_raw"]
    external_smoothed = data["external_smoothed"]
    selected_genes = data["selected_genes"]
    external_feature_audit = data["external_feature_audit"]
    final_training_audit = data["final_training_audit"]

    final_model_summary = pd.DataFrame(
        [
            ["Final classifier", metadata.get("model_name", "linear_svm")],
            ["Feature mode", metadata.get("feature_mode", "stable_genes_top20")],
            ["Gene selection setting", f"Top {metadata.get('top_n_genes_per_region', 20)} genes per region"],
            ["Selected recurrent genes", metadata.get("selected_gene_count", len(selected_genes))],
            ["Spatial neighborhood", f"k = {metadata.get('neighbor_k', 10)} nearest spots"],
            ["Training sections", "12 original labelled Visium DLPFC sections"],
            ["Training spots", metadata.get("training_spot_count", "")],
            ["Training features", metadata.get("training_feature_count", "")],
            ["Target labels", ", ".join(metadata.get("classes", LABEL_ORDER))],
            ["Final saved model", "outputs/models/final_7class_model/final_7class_linear_svm_model.joblib"],
        ],
        columns=["Item", "Value"],
    )

    model_stage_performance = pd.DataFrame(
        [
            [
                "A",
                "Expression-only stable top-20 baseline",
                "No",
                "No",
                baseline["accuracy"],
                baseline["weighted_f1"],
                baseline["macro_f1"],
                "This is the fair no-neighborhood top-20 baseline.",
            ],
            [
                "B",
                "Stable top-20 + k=10 neighborhood features",
                "Yes",
                "No",
                internal_raw["accuracy"],
                internal_raw["weighted_f1"],
                internal_raw["macro_f1"],
                "This estimates the effect of adding neighborhood gene-expression features.",
            ],
            [
                "C",
                "Stable top-20 + k=10 neighborhood features + smoothing",
                "Yes",
                "Yes",
                internal_smoothed["accuracy"],
                internal_smoothed["weighted_f1"],
                internal_smoothed["macro_f1"],
                "This is the best internal setting, but it includes both neighborhood features and post-prediction smoothing.",
            ],
        ],
        columns=[
            "Stage",
            "Setup",
            "Uses neighbor gene features",
            "Uses prediction smoothing",
            "Accuracy",
            "Weighted F1",
            "Macro F1",
            "Strict interpretation",
        ],
    )

    neighborhood_effect_breakdown = pd.DataFrame(
        [
            [
                "No-neighborhood baseline to raw neighborhood model",
                difference(internal_raw["accuracy"], baseline["accuracy"]),
                difference(internal_raw["weighted_f1"], baseline["weighted_f1"]),
                difference(internal_raw["macro_f1"], baseline["macro_f1"]),
                "This is the cleaner estimate of the neighborhood-feature contribution.",
            ],
            [
                "Raw neighborhood model to smoothed neighborhood model",
                difference(internal_smoothed["accuracy"], internal_raw["accuracy"]),
                difference(internal_smoothed["weighted_f1"], internal_raw["weighted_f1"]),
                difference(internal_smoothed["macro_f1"], internal_raw["macro_f1"]),
                "This is the added effect of post-prediction smoothing.",
            ],
            [
                "No-neighborhood baseline to smoothed neighborhood model",
                difference(internal_smoothed["accuracy"], baseline["accuracy"]),
                difference(internal_smoothed["weighted_f1"], baseline["weighted_f1"]),
                difference(internal_smoothed["macro_f1"], baseline["macro_f1"]),
                "This is the total improvement of the final internal setting, not the isolated effect of one component.",
            ],
        ],
        columns=[
            "Comparison",
            "Accuracy change",
            "Weighted F1 change",
            "Macro F1 change",
            "Strict interpretation",
        ],
    )

    internal_external_overall = pd.DataFrame(
        [
            [
                "Internal leave-one-sample-out",
                "Raw",
                internal_raw["accuracy"],
                internal_raw["weighted_f1"],
                internal_raw["macro_f1"],
                "Internal model-selection/evaluation result on original 12 sections.",
            ],
            [
                "Internal leave-one-sample-out",
                "Smoothed",
                internal_smoothed["accuracy"],
                internal_smoothed["weighted_f1"],
                internal_smoothed["macro_f1"],
                "Best internal result, but includes post-prediction smoothing.",
            ],
            [
                "External spatialDLPFC labelled subset",
                "Raw",
                external_raw["accuracy"],
                external_raw["weighted_f1"],
                external_raw["macro_f1"],
                "External result before smoothing on the limited labelled subset.",
            ],
            [
                "External spatialDLPFC labelled subset",
                "Smoothed",
                external_smoothed["accuracy"],
                external_smoothed["weighted_f1"],
                external_smoothed["macro_f1"],
                "Limited external validation result. This should not be reported as full external validation.",
            ],
        ],
        columns=[
            "Validation source",
            "Prediction type",
            "Accuracy",
            "Weighted F1",
            "Macro F1",
            "Strict interpretation",
        ],
    )

    internal_class = data["internal_smoothed_classwise"][
        ["label", "precision", "recall", "f1-score", "support"]
    ].copy()

    external_class = data["external_smoothed_classwise"][
        ["label", "precision", "recall", "f1-score", "support"]
    ].copy()

    internal_class = internal_class.rename(
        columns={
            "precision": "Internal precision",
            "recall": "Internal recall",
            "f1-score": "Internal F1",
            "support": "Internal support",
        }
    )

    external_class = external_class.rename(
        columns={
            "precision": "External precision",
            "recall": "External recall",
            "f1-score": "External F1",
            "support": "External support",
        }
    )

    classwise = internal_class.merge(
        external_class,
        on="label",
        how="outer",
    )

    classwise["F1 change"] = classwise["External F1"] - classwise["Internal F1"]

    classwise["Strict interpretation"] = classwise["label"].map(
        {
            "Layer1": "Moderate external transfer. External support is from Br6522_ant and Br6522_mid only.",
            "Layer2": "Weakest external transfer. Low support and confusion with Layer1/Layer3 are important limitations.",
            "Layer3": "Moderate external transfer with noticeable drop from internal performance.",
            "Layer4": "Weak external transfer, but the drop is smaller because internal performance was already limited.",
            "Layer5": "Strongest external transfer among cortical layers.",
            "Layer6": "Strong external transfer and relatively stable across internal/external testing.",
            "WM": "Good external transfer, but external support is from Br6522_ant and Br6522_mid only.",
        }
    )

    classwise["label"] = pd.Categorical(classwise["label"], categories=LABEL_ORDER, ordered=True)
    classwise = classwise.sort_values("label").reset_index(drop=True)
    classwise["label"] = classwise["label"].astype(str)

    external_sample_performance = data["external_sample_summary"].copy()
    external_sample_performance = external_sample_performance[
        [
            "sample_id",
            "prediction_type",
            "spot_count",
            "accuracy",
            "weighted_f1",
            "macro_f1",
        ]
    ].copy()

    external_sample_performance["Strict interpretation"] = external_sample_performance.apply(
        lambda row: "Br8667_post lacks Layer1 and WM labels, so its macro score is not directly comparable to samples with all seven labels."
        if str(row["sample_id"]) == "Br8667_post"
        else "This sample contains all seven target labels.",
        axis=1,
    )

    external_label_support = build_external_label_support_table(data["external_classwise_all"])

    missing_genes = sorted(
        set(
            gene.strip()
            for value in external_feature_audit["missing_gene_names"].dropna().astype(str)
            for gene in value.split(",")
            if gene.strip()
        )
    )

    external_feature_availability = pd.DataFrame(
        [
            ["Selected final genes", int(external_feature_audit["selected_genes"].max())],
            ["Minimum genes available in external samples", int(external_feature_audit["available_genes"].min())],
            ["Maximum missing genes in external samples", int(external_feature_audit["missing_genes"].max())],
            ["Missing gene names", ", ".join(missing_genes) if missing_genes else "None"],
            ["Strict reporting note", "This is a minor limitation because 84 of 85 selected genes were available externally."],
        ],
        columns=["Item", "Value"],
    )

    internal_feature_check = pd.DataFrame(
        [
            ["Selected genes", int(final_training_audit["selected_genes"].max())],
            ["Minimum available genes in internal training sections", int(final_training_audit["available_genes"].min())],
            ["Maximum missing genes in internal training sections", int(final_training_audit["missing_genes"].max())],
            ["Training feature count", metadata.get("training_feature_count", "")],
            ["Expected feature count", int(metadata.get("selected_gene_count", len(selected_genes))) * 2],
            ["Strict check", "Passed if training feature count equals selected genes × 2 and missing genes = 0."],
        ],
        columns=["Item", "Value"],
    )

    internal_confusion_summary = get_major_confusions(data["internal_confusion_smoothed"])
    external_confusion_summary = get_major_confusions(data["external_confusion_smoothed"])

    weak_layer_analysis = pd.DataFrame(
        [
            [
                "Layer2",
                classwise.loc[classwise["label"] == "Layer2", "Internal support"].iloc[0],
                classwise.loc[classwise["label"] == "Layer2", "External support"].iloc[0],
                classwise.loc[classwise["label"] == "Layer2", "Internal F1"].iloc[0],
                classwise.loc[classwise["label"] == "Layer2", "External F1"].iloc[0],
                "Layer2 has lower internal support than Layer3 and is externally split between Layer1, Layer2, and Layer3. This supports a combined explanation of lower support and weaker separability, not support alone.",
            ],
            [
                "Layer4",
                classwise.loc[classwise["label"] == "Layer4", "Internal support"].iloc[0],
                classwise.loc[classwise["label"] == "Layer4", "External support"].iloc[0],
                classwise.loc[classwise["label"] == "Layer4", "Internal F1"].iloc[0],
                classwise.loc[classwise["label"] == "Layer4", "External F1"].iloc[0],
                "Layer4 is surrounded by higher-support neighboring classes, especially Layer3 and Layer5. Its errors are concentrated toward these neighboring layers.",
            ],
            [
                "Layer5",
                classwise.loc[classwise["label"] == "Layer5", "Internal support"].iloc[0],
                classwise.loc[classwise["label"] == "Layer5", "External support"].iloc[0],
                classwise.loc[classwise["label"] == "Layer5", "Internal F1"].iloc[0],
                classwise.loc[classwise["label"] == "Layer5", "External F1"].iloc[0],
                "Layer5 shows that being between layers is not sufficient to explain weakness. Layer5 likely benefits from higher support and stronger separability.",
            ],
        ],
        columns=[
            "Layer",
            "Internal support",
            "External support",
            "Internal F1",
            "External F1",
            "Strict interpretation",
        ],
    )

    claim_boundaries = pd.DataFrame(
        [
            [
                "The model strongly solved 7-layer prediction internally.",
                "Mostly supported",
                "Use carefully. It is strong internally, but Layer2 and Layer4 remain weaker than the other classes.",
            ],
            [
                "Neighborhood features improved the model.",
                "Supported",
                "The cleaner comparison is no-neighborhood to raw neighborhood: weighted F1 0.6768 to 0.8196.",
            ],
            [
                "Smoothing improved the model.",
                "Supported but smaller effect",
                "The added gain from raw to smoothed was smaller than the gain from adding neighborhood features.",
            ],
            [
                "The model was fully externally validated.",
                "Not supported",
                "Only three spatialDLPFC samples had manual labels. This is limited external validation.",
            ],
            [
                "The model showed external transfer.",
                "Supported",
                "External smoothed weighted F1 was 0.6789 and macro F1 was 0.6708.",
            ],
            [
                "External result is above average.",
                "Not claimable from current evidence",
                "No external baseline or published benchmark on the same labelled subset was run.",
            ],
            [
                "The model is worth pursuing with more labelled datasets.",
                "Supported as a cautious interpretation",
                "The model transferred moderately despite being trained from one dataset source.",
            ],
        ],
        columns=["Claim", "Support status", "Strict wording guidance"],
    )

    final_interpretation_points = pd.DataFrame(
        [
            [
                "Main success",
                "Spatial neighborhood features substantially improved exact 7-class prediction over the expression-only top-20 stable-gene baseline.",
            ],
            [
                "Main internal result",
                f"Internal smoothed weighted F1 = {format_number(internal_smoothed['weighted_f1'])}; macro F1 = {format_number(internal_smoothed['macro_f1'])}.",
            ],
            [
                "Main external result",
                f"External smoothed weighted F1 = {format_number(external_smoothed['weighted_f1'])}; macro F1 = {format_number(external_smoothed['macro_f1'])}.",
            ],
            [
                "External limitation",
                "The external dataset result is limited because only three spatialDLPFC samples had manual layer labels.",
            ],
            [
                "Class limitation",
                "Layer2 and Layer4 remained the weakest labels in both internal and external testing.",
            ],
            [
                "Future work",
                "More manually labelled Visium DLPFC/cortex datasets are needed before claiming robust external generalization.",
            ],
        ],
        columns=["Topic", "Report-ready wording"],
    )

    return {
        "final_model_summary": final_model_summary,
        "model_stage_performance": model_stage_performance,
        "neighborhood_effect_breakdown": neighborhood_effect_breakdown,
        "internal_external_overall": internal_external_overall,
        "classwise_internal_external": classwise,
        "external_sample_performance": external_sample_performance,
        "external_label_support": external_label_support,
        "external_feature_availability": external_feature_availability,
        "internal_feature_check": internal_feature_check,
        "internal_confusion_summary": internal_confusion_summary,
        "external_confusion_summary": external_confusion_summary,
        "weak_layer_analysis": weak_layer_analysis,
        "claim_boundaries": claim_boundaries,
        "final_interpretation_points": final_interpretation_points,
    }


def save_tables(tables):
    for name, table in tables.items():
        table.to_csv(OUTPUT_DIR / f"{name}.csv", index=False)


def make_model_stage_plot(tables):
    df = tables["model_stage_performance"].copy()
    plot_df = df.set_index("Stage")[["Accuracy", "Weighted F1", "Macro F1"]].astype(float)

    ax = plot_df.plot(kind="bar", figsize=(9, 5))
    ax.set_ylim(0, 1)
    ax.set_ylabel("Score")
    ax.set_title("Model Stage Performance")
    ax.legend(loc="lower right")
    plt.xticks(rotation=0)
    plt.tight_layout()

    path = OUTPUT_DIR / "model_stage_performance.png"
    plt.savefig(path, dpi=300)
    plt.close()

    return path


def make_internal_external_plot(tables):
    df = tables["internal_external_overall"].copy()
    df["Group"] = df["Validation source"] + " (" + df["Prediction type"] + ")"
    plot_df = df.set_index("Group")[["Accuracy", "Weighted F1", "Macro F1"]].astype(float)

    ax = plot_df.plot(kind="bar", figsize=(11, 5))
    ax.set_ylim(0, 1)
    ax.set_ylabel("Score")
    ax.set_title("Internal and External Performance")
    ax.legend(loc="lower right")
    plt.xticks(rotation=25, ha="right")
    plt.tight_layout()

    path = OUTPUT_DIR / "internal_external_performance.png"
    plt.savefig(path, dpi=300)
    plt.close()

    return path


def make_classwise_plot(tables):
    df = tables["classwise_internal_external"].copy()
    plot_df = df.set_index("label")[["Internal F1", "External F1"]].astype(float)

    ax = plot_df.plot(kind="bar", figsize=(10, 5))
    ax.set_ylim(0, 1)
    ax.set_ylabel("F1-score")
    ax.set_title("Class-wise Internal vs External F1")
    ax.legend(loc="lower right")
    plt.xticks(rotation=0)
    plt.tight_layout()

    path = OUTPUT_DIR / "classwise_internal_external_f1.png"
    plt.savefig(path, dpi=300)
    plt.close()

    return path


def make_external_sample_plot(tables):
    df = tables["external_sample_performance"].copy()
    df = df[df["prediction_type"] == "smoothed"].copy()

    required_columns = [
        "sample_id",
        "accuracy",
        "weighted_f1",
        "macro_f1",
    ]

    missing_columns = [
        column for column in required_columns
        if column not in df.columns
    ]

    if missing_columns:
        raise KeyError(
            f"Missing required columns for external sample plot: {missing_columns}. "
            f"Available columns: {df.columns.tolist()}"
        )

    plot_df = df.set_index("sample_id")[
        [
            "accuracy",
            "weighted_f1",
            "macro_f1",
        ]
    ].astype(float)

    plot_df = plot_df.rename(
        columns={
            "accuracy": "Accuracy",
            "weighted_f1": "Weighted F1",
            "macro_f1": "Macro F1",
        }
    )

    ax = plot_df.plot(kind="bar", figsize=(9, 5))
    ax.set_ylim(0, 1)
    ax.set_ylabel("Score")
    ax.set_title("External Performance by Labelled spatialDLPFC Sample")
    ax.legend(loc="lower right")
    plt.xticks(rotation=0)
    plt.tight_layout()

    path = OUTPUT_DIR / "external_sample_performance.png"
    plt.savefig(path, dpi=300)
    plt.close()

    return path


def set_document_style(document):
    styles = document.styles

    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10)

    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        if style_name in styles:
            styles[style_name].font.name = "Calibri"


def add_title(document, text):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(18)


def add_small_note(document, text):
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.italic = True
    run.font.size = Pt(9)


def add_table(document, dataframe, font_size=8):
    dataframe = dataframe.copy()

    table = document.add_table(rows=1, cols=len(dataframe.columns))
    table.style = "Table Grid"
    table.autofit = True

    header_cells = table.rows[0].cells

    for index, column in enumerate(dataframe.columns):
        header_cells[index].text = str(column)

    for _, row in dataframe.iterrows():
        cells = table.add_row().cells

        for index, column in enumerate(dataframe.columns):
            value = row[column]
            cells[index].text = format_cell(value)

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)

    document.add_paragraph()


def add_image(document, path, width=6.2):
    path = Path(path)

    if path.exists():
        document.add_picture(str(path), width=Inches(width))
        document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_bullets(document, items):
    for item in items:
        document.add_paragraph(str(item), style="List Bullet")


def add_numbered_list(document, items):
    for item in items:
        document.add_paragraph(str(item), style="List Number")


def create_docx(tables, plots):
    document = Document()
    set_document_style(document)

    add_title(document, "Strict Final 7-Class Visium DLPFC Model Report Package")

    document.add_paragraph(
        "This document summarizes the final model outputs using strict reporting boundaries. It separates neighborhood feature effects from post-prediction smoothing, distinguishes internal validation from limited external validation, and states what can and cannot be claimed."
    )

    document.add_heading("1. Final Model Summary", level=1)
    document.add_paragraph(
        "The final model was trained on the original 12 labelled Visium DLPFC sections. It predicts seven labels: Layer1, Layer2, Layer3, Layer4, Layer5, Layer6, and WM."
    )
    add_table(document, tables["final_model_summary"])

    document.add_heading("2. Model Stage Performance", level=1)
    document.add_paragraph(
        "The final internal result should not be described as a pure neighborhood-only effect, because the best setting includes both neighborhood features and post-prediction smoothing. The stricter breakdown is shown below."
    )
    add_table(document, tables["model_stage_performance"], font_size=7)
    add_image(document, plots["model_stage"])

    document.add_heading("3. Strict Neighborhood Effect Breakdown", level=1)
    document.add_paragraph(
        "The cleanest estimate of the neighborhood-feature contribution is the comparison between the no-neighborhood baseline and the raw neighborhood model. The smoothed result adds a second post-processing effect."
    )
    add_table(document, tables["neighborhood_effect_breakdown"], font_size=7)

    document.add_heading("4. Internal vs External Performance", level=1)
    document.add_paragraph(
        "Internal validation used leave-one-sample-out evaluation across the original 12 labelled Visium DLPFC sections. External testing used only the manually labelled subset of spatialDLPFC. The external result should be reported as limited external validation."
    )
    add_table(document, tables["internal_external_overall"], font_size=7)
    add_image(document, plots["internal_external"])

    document.add_heading("5. Class-wise Internal vs External Comparison", level=1)
    document.add_paragraph(
        "The class-wise result shows that Layer5, Layer6, and WM transferred better externally, while Layer2 and Layer4 remained the weakest labels. Layer1 and WM external support came only from Br6522_ant and Br6522_mid."
    )
    classwise_display = tables["classwise_internal_external"][
        [
            "label",
            "Internal F1",
            "External F1",
            "F1 change",
            "Internal support",
            "External support",
            "Strict interpretation",
        ]
    ].copy()
    add_table(document, classwise_display, font_size=7)
    add_image(document, plots["classwise"])

    document.add_heading("6. External Label Coverage", level=1)
    document.add_paragraph(
        "The external validation subset is incomplete. Br8667_post does not contain Layer1 or WM labels, so per-class external results for Layer1 and WM come from the two Br6522 samples only."
    )
    add_table(document, tables["external_label_support"], font_size=8)

    document.add_heading("7. External Sample-level Results", level=1)
    document.add_paragraph(
        "The external result varies by sample. Br8667_post is less directly comparable because it lacks Layer1 and WM labels."
    )
    add_table(document, tables["external_sample_performance"], font_size=7)
    add_image(document, plots["external_sample"])

    document.add_heading("8. Feature Availability Checks", level=1)
    document.add_paragraph(
        "The internal final model used all selected genes successfully. External validation matched 84 of the 85 selected genes."
    )
    add_table(document, tables["internal_feature_check"], font_size=8)
    add_table(document, tables["external_feature_availability"], font_size=8)

    document.add_heading("9. Confusion Pattern Summary", level=1)
    document.add_paragraph(
        "The confusion patterns indicate that model errors are not random. The most important recurring problems are Layer2 confusion with Layer1/Layer3 and Layer4 confusion with Layer3/Layer5."
    )
    document.add_heading("9.1 Internal smoothed confusion summary", level=2)
    add_table(document, tables["internal_confusion_summary"], font_size=7)
    document.add_heading("9.2 External smoothed confusion summary", level=2)
    add_table(document, tables["external_confusion_summary"], font_size=7)

    document.add_heading("10. Weak-layer Analysis", level=1)
    document.add_paragraph(
        "Layer2 and Layer4 should not be explained only by anatomical position. Layer5 is also between layers but performs much better. The stricter interpretation is that Layer2 and Layer4 combine lower support, neighboring-layer confusion, and weaker separability."
    )
    add_table(document, tables["weak_layer_analysis"], font_size=7)

    document.add_heading("11. Claim Boundaries", level=1)
    document.add_paragraph(
        "This section states which claims are supported and which claims should be avoided."
    )
    add_table(document, tables["claim_boundaries"], font_size=7)

    document.add_heading("12. Final Report-ready Interpretation", level=1)
    add_table(document, tables["final_interpretation_points"], font_size=8)

    document.add_paragraph(
        "Recommended final wording: The final neighborhood-aware Linear SVM model using stable top-20 recurrent genes achieved strong internal performance across the original 12 labelled Visium DLPFC sections. Compared with the expression-only top-20 stable-gene model, adding k=10 neighborhood features improved weighted F1 from 0.6768 to 0.8196, while post-prediction smoothing further improved weighted F1 to 0.8369. On the limited external spatialDLPFC labelled subset, the model achieved moderate external transfer, with weighted F1 of 0.6789 and macro F1 of 0.6708. The model transferred best for Layer5, Layer6, and WM, while Layer2 and Layer4 remained the weakest labels. These results support meaningful but incomplete transfer of DLPFC layer-related signal, and broader validation with more manually labelled external samples is needed before claiming robust generalization."
    )

    output_path = OUTPUT_DIR / "final_model_report_package_strict_detailed.docx"
    document.save(output_path)

    return output_path


def create_markdown(tables):
    path = OUTPUT_DIR / "final_model_report_summary_strict_detailed.md"

    sections = [
        ("Final model summary", tables["final_model_summary"]),
        ("Model stage performance", tables["model_stage_performance"]),
        ("Neighborhood effect breakdown", tables["neighborhood_effect_breakdown"]),
        ("Internal vs external performance", tables["internal_external_overall"]),
        ("Class-wise internal vs external comparison", tables["classwise_internal_external"]),
        ("External label support", tables["external_label_support"]),
        ("External sample performance", tables["external_sample_performance"]),
        ("Feature availability", tables["external_feature_availability"]),
        ("Internal confusion summary", tables["internal_confusion_summary"]),
        ("External confusion summary", tables["external_confusion_summary"]),
        ("Weak-layer analysis", tables["weak_layer_analysis"]),
        ("Claim boundaries", tables["claim_boundaries"]),
        ("Final interpretation points", tables["final_interpretation_points"]),
    ]

    lines = []
    lines.append("# Strict Final 7-Class Visium DLPFC Model Report Package")
    lines.append("")
    lines.append("This summary uses strict reporting boundaries and avoids overclaiming external validation.")
    lines.append("")

    for title, table in sections:
        lines.append(f"## {title}")
        lines.append("")
        lines.append(dataframe_to_markdown(table))
        lines.append("")

    lines.append("## Final recommended wording")
    lines.append("")
    lines.append(
        "The final neighborhood-aware Linear SVM model using stable top-20 recurrent genes achieved strong internal performance across the original 12 labelled Visium DLPFC sections. Compared with the expression-only top-20 stable-gene model, adding k=10 neighborhood features improved weighted F1 from 0.6768 to 0.8196, while post-prediction smoothing further improved weighted F1 to 0.8369. On the limited external spatialDLPFC labelled subset, the model achieved moderate external transfer, with weighted F1 of 0.6789 and macro F1 of 0.6708. The model transferred best for Layer5, Layer6, and WM, while Layer2 and Layer4 remained the weakest labels. These results support meaningful but incomplete transfer of DLPFC layer-related signal, and broader validation with more manually labelled external samples is needed before claiming robust generalization."
    )

    path.write_text("\n".join(lines), encoding="utf-8")

    return path


def create_plain_text_report(tables):
    path = OUTPUT_DIR / "final_model_report_plain_text_strict.txt"

    lines = []
    lines.append("STRICT FINAL MODEL REPORT SUMMARY")
    lines.append("=" * 80)
    lines.append("")
    lines.append("Main conclusion:")
    lines.append(
        "The model achieved strong internal validation and moderate limited external validation. It should not be described as fully externally validated."
    )
    lines.append("")
    lines.append("Most important numbers:")
    lines.append("- Expression-only top-20 stable-gene weighted F1: 0.6768")
    lines.append("- Raw neighborhood internal weighted F1: 0.8196")
    lines.append("- Smoothed neighborhood internal weighted F1: 0.8369")
    lines.append("- External smoothed weighted F1: 0.6789")
    lines.append("- External smoothed macro F1: 0.6708")
    lines.append("")
    lines.append("Strict claim boundaries:")
    lines.append("- Supported: neighborhood features improved internal 7-class prediction.")
    lines.append("- Supported: smoothing added a smaller extra improvement.")
    lines.append("- Supported: limited external transfer exists.")
    lines.append("- Not supported: full external validation across spatialDLPFC.")
    lines.append("- Not supported: above-average external performance without an external baseline.")
    lines.append("")
    lines.append("Weak-layer finding:")
    lines.append(
        "Layer2 and Layer4 remain the weakest labels in both internal and external testing. The most defensible explanation is not anatomical position alone, but lower support, neighboring-layer confusion, and weaker separability."
    )

    path.write_text("\n".join(lines), encoding="utf-8")

    return path


def main():
    data = load_inputs()
    tables = build_summary_tables(data)

    save_tables(tables)

    plots = {
        "model_stage": make_model_stage_plot(tables),
        "internal_external": make_internal_external_plot(tables),
        "classwise": make_classwise_plot(tables),
        "external_sample": make_external_sample_plot(tables),
    }

    docx_path = create_docx(tables, plots)
    markdown_path = create_markdown(tables)
    text_path = create_plain_text_report(tables)

    print("Strict detailed report package created.")
    print("Folder:", OUTPUT_DIR)
    print("DOCX:", docx_path)
    print("Markdown:", markdown_path)
    print("Plain text:", text_path)
    print("CSV tables and PNG figures were also saved.")


if __name__ == "__main__":
    main()